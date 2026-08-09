import io
import re
from typing import Any

from cleaner import clean_text
from extractor import extract_pdf

SECTION_HEADER_SPECS: list[dict[str, Any]] = [
    {
        "type": "summary",
        "keys": [
            "professional summary",
            "career summary",
            "executive summary",
            "about me",
            "summary",
            "profile",
            "objective",
        ],
    },
    {
        "type": "experience",
        "keys": [
            "professional experience",
            "work experience",
            "employment history",
            "career history",
            "work history",
            "employment",
            "internships",
            "internship experience",
            "technical experience",
            "relevant experience",
            "experience",
        ],
    },
    {
        "type": "education",
        "keys": [
            "academic background",
            "academic qualifications",
            "education",
            "academic",
            "qualifications",
        ],
    },
    {
        "type": "skills",
        "keys": [
            "areas of expertise",
            "core competencies",
            "technical skills",
            "key skills",
            "skills",
            "expertise",
            "competencies",
        ],
    },
    {
        "type": "languages",
        "keys": ["language proficiency", "languages", "language"],
    },
    {
        "type": "certifications",
        "keys": ["certifications and courses", "courses", "training", "certifications", "certificates", "licenses"],
    },
    {
        "type": "projects",
        "keys": [
            "personal projects",
            "technical projects",
            "academic projects",
            "key projects",
            "projects",
        ],
    },
    {
        "type": "awards",
        "keys": ["achievements", "honors", "honours", "awards"],
    },
    {
        "type": "volunteer",
        "keys": ["volunteer experience", "volunteering", "memberships", "organisations", "organizations"],
    },
    {
        "type": "interests",
        "keys": ["interests", "hobbies"],
    },
    {
        "type": "references",
        "keys": ["references"],
    },
]

EXACT_ONLY_HEADER_KEYS = {
    "profile",
    "experience",
    "competencies",
    "expertise",
    "skills",
    "education",
    "employment",
    "objective",
    "training",
    "courses",
    "projects",
    "languages",
    "language",
    "interests",
    "references",
    "awards",
    "academic",
    "qualifications",
    "internships",
    "certificates",
    "certifications",
    "licenses",
    "hobbies",
    "summary",
}

EMAIL_RE = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.IGNORECASE)
PHONE_RE = re.compile(
    r"(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{2,4}\)?[\s.-]?)?\d{3}[\s.-]?\d{4}(?:\s*(?:ext\.?|x)\s*\d+)?",
    re.IGNORECASE,
)
URL_RE = re.compile(
    r"(?:https?://)?(?:www\.)?(?:linkedin\.com/\S+|github\.com/\S+)",
    re.IGNORECASE,
)

ALLOWED_RESUME_EXTENSIONS = {".pdf", ".docx"}
ALLOWED_RESUME_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _normalize_heading(line: str) -> str:
    return re.sub(r"[^a-z0-9\s]", "", line.strip().lower()).strip()


def detect_section_type(line: str) -> str | None:
    normalized = _normalize_heading(line)
    if not normalized:
        return None

    for spec in SECTION_HEADER_SPECS:
        for key in sorted(spec["keys"], key=len, reverse=True):
            if key in EXACT_ONLY_HEADER_KEYS:
                if normalized == key:
                    return spec["type"]
            elif normalized == key or normalized.startswith(f"{key} "):
                return spec["type"]

    return None


def _looks_like_contact_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if EMAIL_RE.search(stripped):
        return True
    if URL_RE.search(stripped):
        return True
    if PHONE_RE.search(stripped) and sum(char.isdigit() for char in stripped) >= 7:
        return True
    return False


def _split_skills_items(text: str) -> list[str]:
    if not text.strip():
        return []

    parts = re.split(r"[,|\n•]+", text)
    items = [part.strip() for part in parts if part.strip()]

    if len(items) == 1:
        items = [
            token.strip()
            for token in re.split(r"\s{2,}|\s+and\s+|(?<=[a-z])\s+(?=[A-Z])", text)
            if token.strip() and 1 < len(token.strip()) < 120
        ]

    return items


def _build_line_map(text: str) -> list[dict[str, Any]]:
    lines: list[dict[str, Any]] = []
    offset = 0

    for index, line in enumerate(text.split("\n")):
        lines.append(
            {
                "line_number": index + 1,
                "text": line,
                "char_start": offset,
                "char_end": offset + len(line),
                "section_type": None,
            }
        )
        offset += len(line) + 1

    return lines


def parse_structured_sections(full_text: str) -> dict[str, Any]:
    text = clean_text(full_text)
    raw_lines = text.split("\n")
    line_map = _build_line_map(text)

    sections: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    contact_lines: list[str] = []
    header_lines: list[str] = []
    first_section_index: int | None = None

    for index, line in enumerate(raw_lines):
        section_type = detect_section_type(line)

        if section_type:
            if first_section_index is None:
                first_section_index = index
            if current:
                sections.append(current)
            current = {"type": section_type, "heading": line.strip(), "lines": []}
            if index < len(line_map):
                line_map[index]["section_type"] = section_type
            continue

        if current:
            current["lines"].append(line)
            if index < len(line_map):
                line_map[index]["section_type"] = current["type"]
        elif first_section_index is None:
            if _looks_like_contact_line(line):
                contact_lines.append(line)
                if index < len(line_map):
                    line_map[index]["section_type"] = "contact"
            else:
                header_lines.append(line)

    if current:
        sections.append(current)

    if first_section_index is None:
        contact_lines = [line for line in raw_lines if _looks_like_contact_line(line)]
        header_lines = [line for line in raw_lines if line not in contact_lines]

    structured: dict[str, Any] = {
        "contact": {
            "name": header_lines[0].strip() if header_lines else "",
            "headline": header_lines[1].strip() if len(header_lines) > 1 else "",
            "lines": contact_lines,
            "text": "\n".join(contact_lines).strip(),
        },
        "summary": {"text": "", "paragraphs": []},
        "experience": {"text": "", "paragraphs": []},
        "education": {"text": "", "paragraphs": []},
        "skills": {"text": "", "items": []},
        "additional_sections": [],
        "unassigned": {"text": ""},
    }

    for section in sections:
        body = "\n".join(section["lines"]).strip()
        paragraphs = [paragraph.strip() for paragraph in re.split(r"\n\s*\n", body) if paragraph.strip()]
        payload = {
            "heading": section["heading"],
            "text": body,
            "paragraphs": paragraphs or ([body] if body else []),
        }

        section_type = section["type"]
        if section_type == "summary":
            structured["summary"] = payload
        elif section_type == "experience":
            structured["experience"] = payload
        elif section_type == "education":
            structured["education"] = payload
        elif section_type == "skills":
            structured["skills"] = {
                **payload,
                "items": _split_skills_items(body),
            }
        else:
            structured["additional_sections"].append({"type": section_type, **payload})

    if first_section_index is None and header_lines:
        structured["unassigned"]["text"] = "\n".join(header_lines[2:] + raw_lines).strip()

    return {
        "structured_sections": structured,
        "lines": line_map,
    }


def extract_docx(docx_bytes: bytes) -> dict[str, Any]:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX extraction.") from exc

    document = docx.Document(io.BytesIO(docx_bytes))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    table_lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))

    combined = "\n".join(paragraphs + table_lines)
    full_text = clean_text(combined)

    return {
        "pages": 1,
        "page_texts": [{"text": full_text, "page": 1, "source": "docx"}],
        "full_text": full_text,
        "metadata": {
            "title": document.core_properties.title or "",
            "author": document.core_properties.author or "",
            "pages": 1,
            "has_text_layer": True,
            "extraction_mode": "docx",
            "paragraph_count": len(paragraphs),
            "table_row_count": len(table_lines),
        },
    }


def extract_resume_bytes(file_bytes: bytes, filename: str) -> dict[str, Any]:
    if not file_bytes:
        raise ValueError("Uploaded file is empty.")

    extension = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if extension not in ALLOWED_RESUME_EXTENSIONS:
        raise ValueError("Only PDF and DOCX files are supported.")

    if extension == ".pdf":
        if not file_bytes.startswith(b"%PDF"):
            raise ValueError("File is not a valid PDF.")
        extraction = extract_pdf(file_bytes)
    else:
        extraction = extract_docx(file_bytes)

    full_text = extraction.get("full_text", "")
    cleaned_text = clean_text(full_text)
    from ats_normalizer import normalize_resume_extraction

    normalized = normalize_resume_extraction(cleaned_text)

    return {
        "success": True,
        "filename": filename,
        "file_type": extension.lstrip("."),
        "raw_text": normalized["raw_text"],
        "full_text": normalized["normalized_text"],
        "structured_sections": normalized["structured_sections"],
        "lines": normalized["lines"],
        "normalization": normalized["normalization"],
        "pages": extraction.get("pages", 0),
        "page_texts": extraction.get("page_texts", []),
        "metadata": {
            **(extraction.get("metadata") or {}),
            "ats_normalized": normalized["normalization"]["applied"],
            "normalization_method": normalized["normalization"]["method"],
            "normalization_changed": normalized["normalization"]["changed"],
        },
    }
